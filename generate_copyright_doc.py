#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
软件著作权源代码文档生成器
生成符合软著申请要求的60页核心源代码Word文档
"""

import os
import sys
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# 项目根目录
PROJECT_ROOT = r"E:\1代码\系统\RSCD"

# 每页代码行数(不含空行和注释)
LINES_PER_PAGE = 40

# 目标页数
TARGET_PAGES = 80

# 目标代码行数
TARGET_LINES = TARGET_PAGES * LINES_PER_PAGE  # 3200行

# 排除的入口文件
EXCLUDED_FILES = [
    "start_app.py",
    "__main__.py",
]

# 核心源代码文件列表(按软著要求优先级排序: utils/config > data > service > loss/metric > model)
CORE_FILES = [
    # ========== 10分: 配置与工具模块 ==========
    "change3d_docker/paths.py",
    "change3d_docker/model/utils.py",
    "zhuyaogongneng_docker/theme_manager.py",
    "zhuyaogongneng_docker/display.py",
    "zhuyaogongneng_docker/function/theme_utils.py",
    "change3d_api_docker/path_connector.py",

    # ========== 9分: 数据处理模块 ==========
    "change3d_docker/data/dataset.py",
    "change3d_docker/data/transforms.py",
    "zhuyaogongneng_docker/function/import_before_image.py",
    "zhuyaogongneng_docker/function/import_after_image.py",
    "zhuyaogongneng_docker/function/raster/import_module.py",
    "zhuyaogongneng_docker/function/raster/grid.py",

    # ========== 8分: 服务与业务逻辑 ==========
    "change3d_api_docker/main.py",
    "change3d_api_docker/change_detection_model.py",
    "zhuyaogongneng_docker/app.py",
    "zhuyaogongneng_docker/function/change_cd.py",
    "zhuyaogongneng_docker/function/detection_client.py",
    "zhuyaogongneng_docker/function/batch_processing.py",
    "zhuyaogongneng_docker/function/fishnet_fenge.py",
    "zhuyaogongneng_docker/function/image_display.py",
    "zhuyaogongneng_docker/function/raster.py",
    "zhuyaogongneng_docker/function/raster/batch_processor.py",
    "zhuyaogongneng_docker/function/raster/detection.py",
    "change3d_docker/scripts_app/large_image_BCD.py",
    "change3d_docker/scripts_app/large_raster_BCD.py",
    "change3d_docker/scripts_app/batch_image_BCD.py",
    "change3d_docker/scripts_app/batch_raster_BCD.py",

    # ========== 4分: 评估与度量 ==========
    "change3d_docker/utils/metric_tool.py",

    # ========== 3分: 核心模型与算法(保留核心代码确保软著保护价值) ==========
    "change3d_docker/model/x3d.py",
    "change3d_docker/model/change_decoder.py",
    "change3d_docker/model/trainer.py",
]

def is_effective_code_line(line):
    """
    判断是否为有效代码行

    Args:
        line: 代码行内容

    Returns:
        bool: 是否为有效代码行
    """
    stripped = line.strip()
    # 排除空行和单行注释
    if stripped and not stripped.startswith('#'):
        return True
    return False

def read_source_file(file_path):
    """
    读取源代码文件

    Args:
        file_path: 文件路径

    Returns:
        list: 代码行列表,如果读取失败返回None
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.readlines()
    except UnicodeDecodeError:
        try:
            with open(file_path, 'r', encoding='gbk') as f:
                return f.readlines()
        except:
            return None
    except Exception:
        return None

def create_document_header(doc, project_name="RSCD遥感影像变化检测系统"):
    """
    创建文档头部信息(纯代码模式 - 不添加任何标题)

    Args:
        doc: Document对象
        project_name: 项目名称
    """
    # 纯源代码模式: 不添加任何标题、日期、版本信息
    # 仅包含代码行和行号
    pass

def add_code_to_document(doc, file_path, remaining_lines, line_numbers=True, start_line_number=1):
    """
    将代码添加到文档中(纯代码模式 - 不添加文件标题),并返回实际添加的行数

    Args:
        doc: Document对象
        file_path: 源代码文件路径(相对于项目根目录)
        remaining_lines: 剩余可添加的行数
        line_numbers: 是否显示行号
        start_line_number: 起始行号

    Returns:
        tuple: (添加的代码行数, 下一文件的起始行号, 是否应继续添加)
    """
    if remaining_lines <= 0:
        return 0, start_line_number, False

    full_path = os.path.join(PROJECT_ROOT, file_path)

    # 读取源代码
    code_lines = read_source_file(full_path)
    if not code_lines:
        return 0, start_line_number, True

    # 纯代码模式: 不添加文件标题、文件路径等任何标识信息
    # 直接添加代码内容

    # 添加代码段落
    added_effective_lines = 0
    current_line_number = start_line_number

    for idx, line in enumerate(code_lines, start=1):
        if added_effective_lines >= remaining_lines:
            break

        # 判断是否为有效代码行
        if not is_effective_code_line(line):
            # 无效行仍然添加到文档,但不计入有效行数
            pass
        else:
            added_effective_lines += 1

        # 创建代码段落
        code_para = doc.add_paragraph()

        # 设置段落样式 - 更紧凑的排版
        code_para.paragraph_format.left_indent = Inches(0.25)
        code_para.paragraph_format.line_spacing = 0.9
        code_para.paragraph_format.space_before = Pt(0)
        code_para.paragraph_format.space_after = Pt(0)

        # 添加行号(可选)
        if line_numbers:
            line_num_run = code_para.add_run(f"{current_line_number:4d}: ")
            line_num_run.font.name = 'Consolas'
            line_num_run.font.size = Pt(10)
            line_num_run.font.color.rgb = RGBColor(100, 100, 100)
            current_line_number += 1

        # 添加代码内容(保留缩进)
        code_line = line.replace('\t', '    ')
        if len(code_line) > 150:
            code_line = code_line[:147] + '...\n'

        code_run = code_para.add_run(code_line.rstrip())
        code_run.font.name = 'Consolas'
        code_run.font.size = Pt(10)
        code_run.font.color.rgb = RGBColor(0, 0, 0)

    # 不添加空行分隔 - 保持纯代码连续性
    # 判断是否应该继续添加文件
    should_continue = (remaining_lines - added_effective_lines) > 0

    return added_effective_lines, current_line_number, should_continue

def generate_document():
    """
    生成软件著作权源代码文档(纯代码模式 - 精确控制为60页)
    """
    print("=" * 60)
    print("软件著作权源代码文档生成器 (纯代码模式)")
    print("=" * 60)

    # 创建Word文档
    doc = Document()

    # 设置文档默认字体
    doc.styles['Normal'].font.name = 'Consolas'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
    doc.styles['Normal'].font.size = Pt(10)

    # 不添加文档头部 - 纯代码模式

    # 统计信息
    total_files = 0
    total_lines = 0
    remaining_lines = TARGET_LINES
    current_line_number = 1

    # 遍历核心文件并添加到文档
    print(f"\n开始处理源代码文件(目标: {TARGET_PAGES}页, 约{TARGET_LINES}行)...")

    for file_path in CORE_FILES:
        if remaining_lines <= 0:
            print(f"\n已达到目标行数,停止添加文件。")
            break

        full_path = os.path.join(PROJECT_ROOT, file_path)

        # 检查文件是否存在
        if not os.path.exists(full_path):
            print(f"  [跳过] 文件不存在: {file_path}")
            continue

        # 检查是否为排除的入口文件
        file_name = os.path.basename(file_path)
        if file_name in EXCLUDED_FILES:
            print(f"  [跳过] 入口文件: {file_path}")
            continue

        # 添加代码到文档
        lines_added, current_line_number, should_continue = add_code_to_document(
            doc, file_path, remaining_lines, start_line_number=current_line_number
        )

        if lines_added > 0:
            total_files += 1
            total_lines += lines_added
            remaining_lines -= lines_added
            pages_for_file = (lines_added + LINES_PER_PAGE - 1) // LINES_PER_PAGE
            print(f"  [添加] {file_path} ({lines_added} 行, 约 {pages_for_file} 页, 剩余 {remaining_lines} 行)")

        if not should_continue:
            break

    # 不添加文档尾部统计 - 纯代码模式

    # 保存文档
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    output_filename = f"软著源代码文档_RSCD_80页_{timestamp}.docx"
    output_path = os.path.join(PROJECT_ROOT, output_filename)

    doc.save(output_path)

    # 输出统计信息
    actual_pages = (total_lines + LINES_PER_PAGE - 1) // LINES_PER_PAGE

    print("\n" + "=" * 60)
    print("文档生成完成! (纯源代码模式)")
    print(f"输出文件: {output_path}")
    print(f"包含文件: {total_files} 个")
    print(f"代码行数: {total_lines} 行")
    print(f"实际页数: {actual_pages} 页")
    print("\n文档说明:")
    print("  - 仅包含源代码行和行号")
    print("  - 已移除所有标题、文件路径、统计信息等")
    print("  - 符合软件著作权申请最新要求")
    print("=" * 60)

    return output_path

if __name__ == "__main__":
    try:
        output_file = generate_document()
        print(f"\n✓ 成功生成文档: {output_file}")

        # 询问是否打开文档
        if sys.platform.startswith('win'):
            print("\n提示: 可以双击打开生成的Word文档查看")

    except Exception as e:
        print(f"\n✗ 生成文档时出错: {str(e)}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
